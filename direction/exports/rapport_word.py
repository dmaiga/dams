from collections import defaultdict
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from utils.rapport_utils import jours_non_travailles_individuels


def format_date_courte_fr(d):
    jours = [
        "Lundi", "Mardi", "Mercredi",
        "Jeudi", "Vendredi", "Samedi", "Dimanche"
    ]
    return f"{jours[d.weekday()]} {d.strftime('%d/%m/%Y')}"


def generer_rapport_ventes_word(rapport, date_debut, date_fin, fichier):
    document = Document()

    # ===============================
    # En-tête
    # ===============================
    document.add_heading("RAPPORT DES VENTES", level=1)

    p = document.add_paragraph(
        f"Période : du {format_date_courte_fr(date_debut)} "
        f"au {format_date_courte_fr(date_fin)}\n"
        f"Généré le : {format_date_courte_fr(date.today())}"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ===============================
    # Regroupement par agent
    # ===============================
    agents = defaultdict(list)
    for row in rapport:
        agent = f"{row['agent__user__first_name']} {row['agent__user__last_name']}"
        agents[agent].append(row)

    # ===============================
    # Par agent
    # ===============================
    for agent, ventes in agents.items():
        document.add_heading(f"Agent : {agent}", level=2)

        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        hdr[0].text = "Produit"
        hdr[1].text = "Quantité"
        hdr[2].text = "Poids (kg)"

        jours_vendus = set()
        dernier_jour = None

        for v in ventes:
            jour = v["date_vente__date"]
            jours_vendus.add(jour)

            # 🗓️ Date affichée UNE SEULE FOIS (ligne séparatrice)
            if jour != dernier_jour:
                row_date = table.add_row().cells
                row_date[0].merge(row_date[1]).merge(row_date[2])
                p_date = row_date[0].paragraphs[0]
                p_date.text = format_date_courte_fr(jour)
                p_date.bold = True
                p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dernier_jour = jour

            # Ligne produit
            row = table.add_row().cells
            row[0].text = v["detail_distribution__lot__produit__nom"]
            row[1].text = str(v["total_quantite"])
            row[2].text = f"{v['total_kg']:.2f}"

        # ===============================
        # Jours sans activité
        # ===============================
        document.add_paragraph("Jours sans activité :", style="Heading 3")

        jours_absents = jours_non_travailles_individuels(
            date_debut, date_fin, jours_vendus
        )

        if jours_absents:
            for j in jours_absents:
                document.add_paragraph(f"- {j}", style="List Bullet")
        else:
            document.add_paragraph(
                "Aucun jour ouvré sans activité sur la période."
            )

    document.save(fichier)
